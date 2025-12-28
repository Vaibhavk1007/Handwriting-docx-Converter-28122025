from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
import requests
import os
import io
import re
import json
import base64
import numpy as np
from dataclasses import dataclass
from typing import Dict
from bs4 import BeautifulSoup

from paddleocr import PaddleOCR
from PIL import Image

from html2docx import html2docx
from docx import Document
from docx.shared import Pt, Inches
from queue import Queue
from threading import Thread
import cv2
import uuid

job_queue = Queue()

RESULTS = {}

def worker():
    while True:
        job_id, req = job_queue.get()
        try:
            html = process_job(req)
            RESULTS[job_id] = html
        except Exception as e:
            print("🔥 JOB FAILED:", repr(e))
        finally:
            job_queue.task_done()


# =====================================================
# APP SETUP
# =====================================================

app = FastAPI(title="Ink2Doc Backend")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

OPENROUTER_API_KEY = os.environ.get("OPENROUTER_API_KEY")
if not OPENROUTER_API_KEY:
    raise RuntimeError("OPENROUTER_API_KEY not set")

# =====================================================
# OCR INITIALIZATION (SAFE)
# =====================================================

ocr = None

@app.on_event("startup")
def load_ocr():
    global ocr
    ocr = PaddleOCR(
        lang="en",
        use_textline_orientation=True
    )
    Thread(target=worker, daemon=True).start()


# =====================================================
# DATA MODELS
# =====================================================

class InferenceRequest(BaseModel):
    jobId: str
    image_base64: str
    mode: str = "strict"

class ExportRequest(BaseModel):
    html: str
    title: str | None = "document"

@dataclass
class OCRLine:
    text: str
    x: float
    y: float
    w: float
    h: float
    conf: float

# =====================================================
# UTILITIES
# =====================================================

def decode_image(b64: str) -> np.ndarray:
    img_bytes = base64.b64decode(b64)
    img = Image.open(io.BytesIO(img_bytes)).convert("RGB")
    return np.array(img)

def normalize_ocr(ocr_result):
    page = ocr_result[0]
    lines = []

    for text, conf, box in zip(
        page["rec_texts"],
        page["rec_scores"],
        page["rec_boxes"]
    ):
        x1, y1, x2, y2 = box
        lines.append(OCRLine(
            text=text,
            x=x1,
            y=y1,
            w=x2 - x1,
            h=y2 - y1,
            conf=conf
        ))

    return sorted(lines, key=lambda l: (l.y, l.x))

def group_rows(lines, y_threshold=12):
    rows, current = [], []

    for l in lines:
        if not current or abs(l.y - current[-1].y) <= y_threshold:
            current.append(l)
        else:
            rows.append(current)
            current = [l]

    if current:
        rows.append(current)

    return rows


def process_job(req: InferenceRequest):
    img_np = decode_image(req.image_base64)
    page_width = img_np.shape[1]

    ocr_result = ocr.predict(img_np)

    lines = normalize_ocr(ocr_result)
    rows = group_rows(lines)
    rows = build_reading_order(rows)
    rows = normalize_layout(rows)
    tables, rows = detect_tables(rows)
    rows = merge_page_continuations(rows)

    

    rule_labels, unresolved = classify_rows_rule_based(rows, page_width)

    llm_labels = {}
    high_conf = [
        i for i in unresolved
        if row_confidence(rows[i]) >= 0.6
    ]

    if high_conf:
        payload = tokenize_rows([rows[i] for i in high_conf])
        llm_labels = classify_rows_with_llm(payload)

    final_labels = {**rule_labels, **llm_labels}

    html = rows_to_html(rows, final_labels, page_width, tables)
    html = validate_and_repair_html(html)

    return html



def build_reading_order(rows):
    """
    Creates a stable reading order graph:
    - rows top to bottom
    - tokens left to right
    """
    ordered = []
    for row in rows:
        ordered.append(sorted(row, key=lambda l: l.x))
    return ordered    

def normalize_layout(rows):
    """
    Normalizes vertical spacing and removes noise rows.
    """
    normalized = []
    for row in rows:
        clean = [l for l in row if l.conf > 0.4 and l.text.strip()]
        if clean:
            normalized.append(clean)
    return normalized

def tokenize_rows(rows):
    tokens = {}
    for i, row in enumerate(rows):
        tokens[str(i)] = {
            "text": " ".join(l.text for l in row),
            "y": float(row[0].y),
            "avg_conf": float(sum(float(l.conf) for l in row) / len(row)),
            "len": int(len(row)),
        }
    return tokens


def validate_and_repair_html(html: str) -> str:
    soup = BeautifulSoup(html, "html.parser")

    # Ensure document wrapper
    if not soup.find("div", class_="document"):
        wrapper = soup.new_tag("div", **{"class": "document"})
        wrapper.append(soup)
        soup = wrapper

    # Remove empty paragraphs
    for p in soup.find_all("p"):
        if not p.get_text(strip=True):
            p.decompose()

    return str(soup)        


# =====================================================
# RULE-BASED CLASSIFICATION
# =====================================================


def is_header_row(row, page_width):
    left, right = split_columns(row, page_width)
    left_text = " ".join(l.text.lower() for l in left)
    right_text = " ".join(l.text.lower() for l in right)

    return (
        right
        and ("date" in right_text or re.search(r"\d{2}[./-]\d{2}[./-]\d{4}", right_text))
        and len(left) > 0
    )

def classify_rows_rule_based(rows, page_width):
    labels = {}
    unresolved = []

    for i, row in enumerate(rows):
        texts = [l.text.strip() for l in row if l.text.strip()]
        if not texts:
            continue

        joined = " ".join(texts)

        # 🔥 1. HEADER (ADVOCATE + DATE)
        if is_header_row(row, page_width):
            labels[str(i)] = "doc-header"
        
        elif joined.isupper() and 15 < len(joined) < 60:
            labels[str(i)] = "doc-title"

        # 2. PAGE CONTINUATION
        elif joined.startswith("..."):
            labels[str(i)] = "page-continuation"

        # 3. SALUTATION
        elif texts[0].lower() == "to":
            labels[str(i)] = "doc-salutation"

        # 4. RECIPIENT META BLOCK
        elif texts[0].startswith(("1.", "2.", "3.")) and "residing" in joined.lower():
            labels[str(i)] = "doc-meta"

        # 5. BODY
        elif (
            "instructions of my clients" in joined.lower()
            or re.match(r"^\d+\.", texts[0])
        ):
            labels[str(i)] = "doc-body"

        else:
            unresolved.append(i)

    return labels, unresolved

# =====================================================
# LLM PROMPTS (TWO MODES)
# =====================================================

SYSTEM_PROMPT_STRICT = """
You are a document row classification engine.

Task:
Classify each row into its correct document section.

Rules:
- Do NOT rewrite text
- Do NOT normalize spelling
- Do NOT merge lines
- Do NOT infer meaning

Return ONLY valid JSON.

Format:
{
  "labels": {
    "3": "doc-meta",
    "4": "doc-body"
  }
}

Allowed labels:
- doc-title
- doc-meta
- doc-subject
- doc-salutation
- doc-body
- doc-signoff
- page-continuation
"""

SYSTEM_PROMPT_NORMAL = """
You are a document formatting assistant.

Input:
HTML of a document that is already structurally correct.

You MAY:
- Improve spacing
- Improve paragraph flow
- Split very long paragraphs

You MUST NOT:
- Change meaning
- Remove content
- Add content
- Reorder sections
- Alter numbering or dates

Return ONLY valid HTML.
"""

# =====================================================
# LLM HELPERS
# =====================================================


def extract_json_safe(text: str) -> dict:
    if not text:
        raise ValueError("Empty LLM response")

    # Remove markdown fences if any
    text = re.sub(r"```json|```", "", text, flags=re.IGNORECASE).strip()

    match = re.search(r"\{[\s\S]*\}", text)
    if not match:
        raise ValueError(f"No JSON found in LLM response:\n{text}")

    return json.loads(match.group(0))


def safe_extract_labels(text: str) -> Dict[str, str]:
    match = re.search(r"\{[\s\S]*\}", text)
    if not match:
        raise ValueError("No JSON found in LLM response")
    return json.loads(match.group(0))["labels"]

def classify_rows_with_llm(payload):
    try:
        res = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                "Content-Type": "application/json",
            },
            json={
                "model": "mistralai/mistral-7b-instruct",
                "temperature": 0.0,
                "messages": [
                    {"role": "system", "content": SYSTEM_PROMPT_STRICT},
                    {
                        # 🔑 IMPORTANT: wrap like notebook
                        "role": "user",
                        "content": json.dumps({"rows": payload}),
                    },
                ],
            },
            timeout=60,
        )

        res.raise_for_status()

        content = res.json()["choices"][0]["message"]["content"]

        if not content or not content.strip():
            print("⚠️ LLM returned empty content, skipping LLM")
            return {}

        parsed = extract_json_safe(content)
        return parsed.get("labels", {})

    except Exception as e:
        print("⚠️ LLM classification failed, falling back to rules:", repr(e))
        return {}


def detect_tables(rows, x_tolerance=15):
    """
    Detect tables using grid clustering.
    Returns:
      tables: list of {start, end, cols}
      non_table_rows: remaining rows
    """
    tables = []
    used = set()

    for i in range(len(rows) - 1):
        if i in used:
            continue

        row = rows[i]
        if len(row) < 2:
            continue

        # detect aligned columns
        x_positions = [l.x for l in row]

        matched = [i]
        for j in range(i + 1, len(rows)):
            other = rows[j]
            if len(other) != len(row):
                break

            if all(
                abs(other[k].x - x_positions[k]) < x_tolerance
                for k in range(len(row))
            ):
                matched.append(j)
            else:
                break

        if len(matched) >= 2:
            tables.append({
                "start": matched[0],
                "end": matched[-1],
                "cols": len(row),
            })
            used.update(matched)

    non_table_rows = [
        r for idx, r in enumerate(rows) if idx not in used
    ]

    return tables, non_table_rows


# =====================================================
# HTML RENDERING
# =====================================================


def split_columns(row, page_width, ratio=0.55):
    """
    Split a row into left and right columns based on x position.
    """
    left, right = [], []
    for l in row:
        if l.x < page_width * ratio:
            left.append(l)
        else:
            right.append(l)
    return left, right

def render_table(table_rows):
    html = ['<table border="1" style="border-collapse:collapse; width:100%;">']
    for row in table_rows:
        html.append("<tr>")
        for cell in row:
            html.append(f"<td>{cell.text}</td>")
        html.append("</tr>")
    html.append("</table>")
    return "\n".join(html)

def row_confidence(row):
    return sum(l.conf for l in row) / len(row)

def merge_page_continuations(rows):
    merged = []
    buffer = None

    for row in rows:
        text = " ".join(l.text for l in row)

        if buffer and text and not text[0].isupper():
            buffer.extend(row)
        else:
            if buffer:
                merged.append(buffer)
            buffer = row

    if buffer:
        merged.append(buffer)

    return merged

def is_legal_continuation(text: str) -> bool:
    text = text.lower()
    return any(
        key in text
        for key in ["s/o", "d/o", "w/o", "aged about", "residing", "executed"]
    )    

def should_merge(prev: str, curr: str) -> bool:
    curr_l = curr.lower()

    return (
        curr_l.startswith(("1.", "2.", "3."))
        or any(
            key in curr_l
            for key in [
                "s/o", "d/o", "w/o",
                "aged about",
                "residing",
                "executed",
                "no.",
            ]
        )
    )

BLOCK_LABELS = {
    "doc-header",
    "doc-title",
    "doc-salutation",
    "doc-meta",
    "page-continuation",
}


def rows_to_html(rows, labels, page_width, tables):
    html = ['<div class="document">']
    buffer = ""

    # Build quick lookup for table starts
    table_starts = {
        t["start"]: t for t in tables
    }

    skip_until = -1

    for i, row in enumerate(rows):
        # ⛔ Skip rows that belong to a table already rendered
        if i <= skip_until:
            continue

        # 🟦 TABLE START
        if i in table_starts:
            table = table_starts[i]
            table_rows = rows[table["start"]: table["end"] + 1]

            html.append(render_table(table_rows))

            skip_until = table["end"]
            continue

        # 🟩 NORMAL ROW FLOW (your existing logic)
        label = labels.get(str(i), "doc-body")
        left, right = split_columns(row, page_width)

        # PAGE CONTINUATION
        if label == "page-continuation":
            html.append(
                f'<p class="doc-meta" style="text-align:right;">{" ".join(l.text for l in row)}</p>'
            )
            continue

        # HEADER ROW (left + right)
        if label == "doc-header":
            html.append(
                '<div class="doc-header" style="display:flex; justify-content:space-between; margin-bottom:12px;">'
            )

            # LEFT: Advocate / address
            html.append('<div>')
            for l in left:
                html.append(l.text + "<br>")
            html.append('</div>')

            # RIGHT: Date
            html.append('<div style="text-align:right;">')
            for r in right:
                html.append(r.text + "<br>")
            html.append('</div>')

            html.append('</div>')
            continue

        # TITLE
        if label == "doc-title":
            html.append(
                f'<p class="doc-title" '
                f'style="text-align:center; font-weight:bold; margin:12px 0;">'
                f'{" ".join(l.text for l in left)}'
                f'</p>'
            )
            continue

        # NORMAL BODY
        paragraph = " ".join(l.text for l in left)

        if buffer:
            if should_merge(buffer, paragraph):
                buffer += " " + paragraph
                continue
            else:
                html.append(f'<p class="doc-body">{buffer}</p>')
                buffer = ""

        if label == "doc-body":
            buffer = paragraph
    
    if buffer:
        html.append(f'<p class="doc-body">{buffer}</p>')

    html.append('</div>')
    return "\n".join(html)




def render_debug(img, rows, labels):
    debug = img.copy()

    for i, row in enumerate(rows):
        label = labels.get(str(i), "")
        for l in row:
            cv2.rectangle(
                debug,
                (int(l.x), int(l.y)),
                (int(l.x + l.w), int(l.y + l.h)),
                (0, 255, 0),
                1,
            )

        if row:
            cv2.putText(
                debug,
                label,
                (int(row[0].x), max(10, int(row[0].y - 5))),
                cv2.FONT_HERSHEY_SIMPLEX,
                0.4,
                (255, 0, 0),
                1,
            )

    return debug

# =====================================================
# PROCESS (STRICT MODE ALWAYS)
# =====================================================

@app.post("/process")
def enqueue(req: InferenceRequest):
    job_queue.put((req.jobId, req))
    return {"jobId": req.jobId, "status": "queued"}

@app.post("/debug")
def debug(req: InferenceRequest):
    try:
        img_np = decode_image(req.image_base64)

        ocr_result = ocr.predict(img_np)
        lines = normalize_ocr(ocr_result)
        rows = group_rows(lines)
        tables, rows = detect_tables(rows)
        rows = build_reading_order(rows)
        rows = normalize_layout(rows)
        rows = merge_page_continuations(rows)
        page_width = img_np.shape[1]
        rule_labels, unresolved = classify_rows_rule_based(rows, page_width)
        debug_img = render_debug(img_np, rows, rule_labels)

        _, buf = cv2.imencode(".png", debug_img)

        return StreamingResponse(
            io.BytesIO(buf.tobytes()),
            media_type="image/png"
        )

    except Exception as e:
        print("🔥 DEBUG ERROR:", repr(e))
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/result/{job_id}")
def get_result(job_id: str):
    html = RESULTS.get(job_id)
    if not html:
        raise HTTPException(404, "Result not ready")
    return {"html": html}

# =====================================================
# POLISH (NORMAL MODE)
# =====================================================

@app.post("/polish")
def polish(req: ExportRequest):
    try:
        res = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                "Content-Type": "application/json",
            },
            json={
                "model": "mistralai/mistral-7b-instruct",
                "temperature": 0.3,
                "messages": [
                    {"role": "system", "content": SYSTEM_PROMPT_NORMAL},
                    {"role": "user", "content": req.html},
                ],
            },
            timeout=60,
        )

        res.raise_for_status()
        html = res.json()["choices"][0]["message"]["content"]

        return {"html": html}

    except Exception as e:
        print("🔥 POLISH ERROR:", repr(e))
        raise HTTPException(status_code=500, detail=str(e))

# =====================================================
# DOCX EXPORT
# =====================================================

def polish_docx(doc: Document):
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)

    for table in doc.tables:
        table.style = "Table Grid"
        table.autofit = False
        width = Inches(6.5) / len(table.columns)
        for row in table.rows:
            for cell in row.cells:
                cell.width = width

def sanitize_filename(name: str) -> str:
    safe = re.sub(r"[^A-Za-z0-9_\-\.]", "_", name)
    return (safe or "document")[:200]

@app.post("/export")
def export_doc(req: ExportRequest):
    if not req.html.strip():
        raise HTTPException(status_code=400, detail="missing html")

    buf = html2docx(req.html, title=req.title or "document")
    buf.seek(0)

    doc = Document(buf)
    polish_docx(doc)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)

    return StreamingResponse(
        out,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{sanitize_filename(req.title)}.docx"'
        },
    )
